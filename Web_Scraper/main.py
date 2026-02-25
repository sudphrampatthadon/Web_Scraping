import cloudscraper
from bs4 import BeautifulSoup
from docx import Document
import time
from urllib.parse import urljoin
import re

doc = Document()

# Create the scraper instance
scraper = cloudscraper.create_scraper()

url = input("Enter the URL of the first chapter: ").strip()
base_url = '' # input Base URL like "https://www.novelhall.com/"

title_fixed = False 
filename = "novel.docx" # Base file name
count = 0 # Counter web scaping 

while url:
    print(f"Scraping: {url}") 
    response = scraper.get(url)

    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        count += 1 
        
        # 1. Handle Filename (Once)
        if not title_fixed:
            title_tag = soup.find('a', id='bookname')
            if title_tag:
                novel_title = title_tag.text.strip()
                clean_title = re.sub(r'[\\/*?:"<>|]', " ", novel_title)
                filename = f"{' '.join(clean_title.split())}.docx"
                title_fixed = True

        # 2. Get Chapter Title
        header_content = soup.find('div', class_='single-header')
        if header_content and header_content.h1:
            doc.add_heading(header_content.h1.get_text(), level=1)

        # 3. Get Chapter Content
        content = soup.find('div', id='htmlContent')
        if content:
            for text in content.stripped_strings:
                doc.add_paragraph(text)

        # 4. Find the "Next" link
        next_page = soup.find('a', string='Next')
        if next_page and 'href' in next_page.attrs:
            url = urljoin(url, next_page['href'])
            
            # Periodic Save
            if count % 10 == 0:
                doc.save(filename)
                print(f"Progress saved at chapter {count}...")
            
            time.sleep(1)
        else:
            print("No more chapters found.")
            url = None  

    elif response.status_code == 403:
        print("FAILED: Cloudflare 403. Try a VPN or Hotspot.")
        break
    else:
        print(f"Failed: Status {response.status_code}")
        break

doc.save(filename)
print(f"Final document saved as '{filename}'.")